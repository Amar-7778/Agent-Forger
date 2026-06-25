import os
import PyPDF2
from docx import Document
import pandas as pd
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document as LangChainDocument

CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "chroma_db")
DOCS_DIR = os.path.join("data", "company_docs")

# Ensure directories exist
os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)
os.makedirs(DOCS_DIR, exist_ok=True)

# Initialize embeddings and vector store
try:
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vector_store = Chroma(persist_directory=CHROMA_PERSIST_DIR, embedding_function=embeddings)
except Exception as e:
    print(f"Failed to initialize embeddings: {e}")
    embeddings = None
    vector_store = None

def extract_text_from_file(file_path: str) -> str:
    ext = file_path.lower().split('.')[-1]
    text = ""
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif ext == 'pdf':
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif ext == 'docx':
            doc = Document(file_path)
            text = " ".join([para.text for para in doc.paragraphs])
        elif ext == 'csv':
            df = pd.read_csv(file_path)
            text = df.to_string()
    except Exception as e:
        print(f"Error parsing {file_path}: {e}")
    return text

def ingest_document(file_name: str, file_bytes: bytes) -> tuple[bool, str]:
    if vector_store is None:
        return False, "Vector store is not initialized."
        
    # Save file temporarily
    file_path = os.path.join(DOCS_DIR, file_name)
    with open(file_path, 'wb') as f:
        f.write(file_bytes)
    
    text = extract_text_from_file(file_path)
    if not text.strip():
        return False, "Could not extract text from file."
        
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(text)
    
    documents = [LangChainDocument(page_content=chunk, metadata={"source": file_name}) for chunk in chunks]
    vector_store.add_documents(documents)
    return True, f"Successfully ingested {len(chunks)} chunks from {file_name}"

def query_knowledge_base(query: str, k: int = 3) -> str:
    if vector_store is None:
        return "Knowledge base is offline."
        
    results = vector_store.similarity_search(query, k=k)
    if not results:
        return "No relevant information found in the knowledge base."
    
    context = "\n\n".join([f"Source ({doc.metadata.get('source', 'Unknown')}):\n{doc.page_content}" for doc in results])
    return context
